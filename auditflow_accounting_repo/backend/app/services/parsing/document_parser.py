from io import BytesIO
from pathlib import Path

import fitz
from docx import Document

from app.models.document_section import DocumentFormat, ParsedDocument
from app.services.parsing.text_normalization import normalize_text


class UnsupportedDocumentFormatError(ValueError):
    pass


class EmptyParsedTextError(ValueError):
    pass


class DocumentParser:
    def parse(self, filename: str, content: bytes) -> ParsedDocument:
        document_format = self._detect_format(filename)
        raw_text = self._extract_text(document_format, content)
        text = normalize_text(raw_text)
        if not text:
            raise EmptyParsedTextError("No extractable text was found in the document.")

        return ParsedDocument(
            filename=Path(filename).name,
            document_format=document_format,
            text=text,
        )

    def _detect_format(self, filename: str) -> DocumentFormat:
        extension = Path(filename).suffix.lower()
        if extension == ".pdf":
            return DocumentFormat.PDF
        if extension == ".docx":
            return DocumentFormat.DOCX
        if extension == ".txt":
            return DocumentFormat.TXT
        raise UnsupportedDocumentFormatError(
            "Supported document formats are PDF, DOCX, and TXT."
        )

    def _extract_text(self, document_format: DocumentFormat, content: bytes) -> str:
        if document_format == DocumentFormat.PDF:
            return self._extract_pdf_text(content)
        if document_format == DocumentFormat.DOCX:
            return self._extract_docx_text(content)
        if document_format == DocumentFormat.TXT:
            return self._extract_txt_text(content)
        raise UnsupportedDocumentFormatError(
            "Supported document formats are PDF, DOCX, and TXT."
        )

    def _extract_txt_text(self, content: bytes) -> str:
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            return content.decode("latin-1")

    def _extract_docx_text(self, content: bytes) -> str:
        document = Document(BytesIO(content))
        parts: list[str] = []

        parts.extend(paragraph.text for paragraph in document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))

        return "\n".join(parts)

    def _extract_pdf_text(self, content: bytes) -> str:
        with fitz.open(stream=content, filetype="pdf") as document:
            return "\n".join(page.get_text("text") for page in document)
