from io import BytesIO

import fitz
import pytest
from docx import Document

from app.models.document_section import DocumentFormat
from app.services.parsing.document_parser import (
    DocumentParser,
    EmptyParsedTextError,
    UnsupportedDocumentFormatError,
)
from app.services.parsing.text_normalization import normalize_text


def test_normalize_text_collapses_spacing_and_blank_lines() -> None:
    text = "  Accounting\tmemo\r\n\r\n\r\n  Posting   logic  "

    assert normalize_text(text) == "Accounting memo\n\nPosting logic"


def test_parse_txt_document() -> None:
    parser = DocumentParser()

    parsed = parser.parse(
        filename="walkthrough.txt",
        content=b"Revenue cycle\r\n\r\nDebit cash   credit revenue",
    )

    assert parsed.filename == "walkthrough.txt"
    assert parsed.document_format == DocumentFormat.TXT
    assert parsed.text == "Revenue cycle\n\nDebit cash credit revenue"


def test_parse_docx_document() -> None:
    parser = DocumentParser()
    document = Document()
    document.add_paragraph("Month-end close")
    document.add_paragraph("Review journal entries before posting.")
    buffer = BytesIO()
    document.save(buffer)

    parsed = parser.parse(filename="memo.docx", content=buffer.getvalue())

    assert parsed.document_format == DocumentFormat.DOCX
    assert "Month-end close" in parsed.text
    assert "Review journal entries before posting." in parsed.text


def test_parse_pdf_document() -> None:
    parser = DocumentParser()
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Chart of accounts policy")
    content = document.tobytes()
    document.close()

    parsed = parser.parse(filename="policy.pdf", content=content)

    assert parsed.document_format == DocumentFormat.PDF
    assert "Chart of accounts policy" in parsed.text


def test_parser_rejects_unsupported_format() -> None:
    parser = DocumentParser()

    with pytest.raises(UnsupportedDocumentFormatError):
        parser.parse(filename="ledger.csv", content=b"account,value")


def test_parser_rejects_empty_extractable_text() -> None:
    parser = DocumentParser()

    with pytest.raises(EmptyParsedTextError):
        parser.parse(filename="empty.txt", content=b"   \n\n")
